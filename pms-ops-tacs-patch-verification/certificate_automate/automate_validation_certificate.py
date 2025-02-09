import os
import re
from collections import defaultdict
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt


class PatchFileManager:
    def __init__(self, base_path):
        self.base_path = base_path

    def get_patch_dir(self):
        subdirs = [
            f
            for f in os.listdir(self.base_path)
            if os.path.isdir(os.path.join(self.base_path, f))
        ]
        if len(subdirs) == 1:
            return subdirs[0]
        else:
            raise Exception("data 폴더 하위에 폴더가 하나만 있어야 합니다.")

    def find_v1_dir(self):
        for root, dirs, _ in os.walk(self.base_path):
            if "pms_patch" in dirs:
                pms_patch_path = os.path.join(root, "pms_patch")
                v1_path = os.path.join(pms_patch_path, "v1")
                if os.path.exists(v1_path):
                    return v1_path
        raise Exception("pms_patch/v1 폴더를 찾을 수 없습니다.")

    def map_ms_patch_titles(self, file_list, ms_excel_file):
        df = pd.read_excel(ms_excel_file)
        ms_patch_title_map = dict(zip(df["패치파일"], df["제목"]))

        mapped_files = []

        for file in file_list:
            title = ms_patch_title_map.get(file, None)
            if title:
                mapped_files.append(f"{title}\n{file}")
            else:
                mapped_files.append(file)

        return mapped_files

    def list_files_in_v1(self, v1_dir_path, ms_excel_file):
        if not os.path.exists(v1_dir_path) or not os.path.isdir(v1_dir_path):
            return "v1 폴더가 존재하지 않습니다."

        result = []
        section_count = 0
        sw_files_grouped = defaultdict(list)

        for root, dirs, files in os.walk(v1_dir_path):
            if "sw_files" in dirs:
                dirs.remove("sw_files")
            relative_path = os.path.relpath(root, v1_dir_path).replace(os.sep, "/")
            filtered_files = [
                file
                for file in files
                if file.endswith((".cab", ".exe")) and not file.endswith(".zip")
            ]
            if filtered_files:
                section_count += 1
                prefix = chr(96 + section_count) + "."
                mapped_files = list(
                    dict.fromkeys(
                        self.map_ms_patch_titles(filtered_files, ms_excel_file)
                    )
                )
                file_list = [f"{i + 1}) {file}" for i, file in enumerate(mapped_files)]
                section = (
                    f"{prefix} {relative_path} 하위 {len(filtered_files)}개 파일 확인\n"
                    + "\n".join(file_list)
                )
                result.append(section)

        sw_files_path = os.path.join(v1_dir_path, "sw_files")
        if os.path.exists(sw_files_path):
            for root, _, files in os.walk(sw_files_path):
                relative_path = os.path.relpath(root, sw_files_path).replace(
                    os.sep, "/"
                )
                base_dir = relative_path.split("/")[0] if relative_path != "." else ""
                for file in files:
                    file_path = os.path.relpath(
                        os.path.join(root, file), os.path.join(sw_files_path, base_dir)
                    ).replace(os.sep, "/")
                    sw_files_grouped[base_dir].append(file_path)

            for base_dir, grouped_files in sw_files_grouped.items():
                section_count += 1
                prefix = chr(96 + section_count) + "."
                mapped_files = list(
                    dict.fromkeys(
                        self.map_ms_patch_titles(grouped_files, ms_excel_file)
                    )
                )
                file_list = [f"{i + 1}) {file}" for i, file in enumerate(mapped_files)]
                section = (
                    f"{prefix} sw_files/{base_dir} 하위 {len(grouped_files)}개 파일 확인\n"
                    + "\n".join(file_list)
                )
                result.append(section)

        return "\n\n".join(result)

    def get_numeric_dirs(self):
        patch_dir = self.get_patch_dir()
        standalone_path = os.path.join(
            self.base_path, patch_dir, "pms_patch", "standalone"
        )

        if not os.path.exists(standalone_path):
            raise FileNotFoundError(
                f"'standalone' 폴더를 찾을 수 없습니다: {standalone_path}"
            )

        numeric_dirs = [
            f
            for f in os.listdir(standalone_path)
            if os.path.isdir(os.path.join(standalone_path, f)) and f.isdigit()
        ]
        numeric_dirs.sort()
        return numeric_dirs


class ExcelDataProcessor:
    def __init__(self):
        pass

    def map_patch_titles(self, file_list, ms_excel_file):
        df = pd.read_excel(ms_excel_file)
        ms_patch_title_map = dict(zip(df["패치파일"], df["제목"]))
        mapped_files = []
        for file in file_list:
            title = ms_patch_title_map.get(file, None)
            if title:
                mapped_files.append(f"{title}\n{file}")
            else:
                mapped_files.append(file)
        return mapped_files

    def count_21h2_titles(self, ms_excel_file):
        df = pd.read_excel(ms_excel_file)
        return df["제목"].str.contains("21H2", na=False).sum()

    def count_22h2_titles(self, ms_excel_file):
        df = pd.read_excel(ms_excel_file)
        return df["제목"].str.contains("22H2", na=False).sum()

    def get_kb_ids_grouped(self, ms_excel_file):
        df = pd.read_excel(ms_excel_file)

        kb_21h2 = (
            df[df["제목"].str.contains("21H2", na=False)]["KB ID"].dropna().unique()
        )
        kb_22h2 = (
            df[df["제목"].str.contains("22H2", na=False)]["KB ID"].dropna().unique()
        )

        kb_21h2 = sorted(set(f"KB{str(kb_id)}" for kb_id in kb_21h2))
        kb_22h2 = sorted(set(f"KB{str(kb_id)}" for kb_id in kb_22h2))

        return ", ".join(kb_21h2), ", ".join(kb_22h2)

    def get_office_kb_ids(self, ms_excel_file):
        df = pd.read_excel(ms_excel_file)
        office_kb_ids = (
            df[df["제목"].str.contains("보안 업데이트", na=False)]["KB ID"]
            .dropna()
            .unique()
        )
        office_kb_ids = sorted(set(f"KB{str(kb_id)}" for kb_id in office_kb_ids))
        return len(office_kb_ids), ", ".join(office_kb_ids)

    def get_software_versions(self, ms_excel_file, software_names):
        df = pd.read_excel(ms_excel_file)
        versions = {}
        for software in software_names:
            matched_versions = (
                df[df["패치명"].str.contains(software, na=False)]["버전"]
                .dropna()
                .unique()
            )
            versions[software] = ", ".join(sorted(set(matched_versions)))
        return versions

    def get_edge_version(self, excel_file):
        df = pd.read_excel(excel_file)
        edge_versions = (
            df[df["제목"].str.contains("Edge", na=False)]["제목"].dropna().tolist()
        )
        version_pattern = r"빌드\s*([\d\.]+)"
        for title in edge_versions:
            match = re.search(version_pattern, title)
            if match:
                return match.group(1)
        return "버전 정보 없음"


class DocumentUpdater:
    def __init__(self, base_path, excel_processor, file_manager):
        self.base_path = base_path
        self.excel_processor = excel_processor
        self.file_manager = file_manager

    def set_font_size(self, document, size):
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(size)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(size)

    def update_completion_date(self, docx_path):
        document = Document(docx_path)
        today = datetime.now()

        last_table = document.tables[-1]
        last_row = last_table.rows[-1]
        last_cell = last_row.cells[-1]

        last_cell.text = last_cell.text.replace("년", f"{today.year}년", 1)
        last_cell.text = last_cell.text.replace("월", f"{today.month:02d}월", 1)

        if "일" in last_cell.text:
            last_occurrence_index = last_cell.text.rfind("일")
            last_cell.text = (
                last_cell.text[:last_occurrence_index]
                + f"{today.day:02d}일"
                + last_cell.text[last_occurrence_index + 1 :]
            )

        document.save(docx_path)

    def update_docx_with_content(
        self, source_file_path, content_to_add, new_file_path, month
    ):
        document = Document(source_file_path)
        try:
            data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
            patch_dir_name = self.get_patch_dir(data_dir)
            current_year = datetime.now().year
            previous_year, previous_month = self.get_previous_year_and_month()
            ms_excel_file = os.path.join(data_dir, "ms_patch_list.xlsx")
            sw_excel_file = os.path.join(data_dir, "sw_patch_list.xlsx")
            count_21h2 = self.excel_processor.count_21h2_titles(ms_excel_file)
            count_22h2 = self.excel_processor.count_22h2_titles(ms_excel_file)
            kb_ids_21h2, kb_ids_22h2 = self.excel_processor.get_kb_ids_grouped(
                ms_excel_file
            )
            office_count, office_kb_ids = self.excel_processor.get_office_kb_ids(
                ms_excel_file
            )

            edge_version = self.excel_processor.get_edge_version(ms_excel_file)

            software_list = ["Adobe Acrobat Reader", "BandiZip", "Chrome"]
            software_versions = self.excel_processor.get_software_versions(
                sw_excel_file, software_list
            )

            table = document.tables[1]
            for row in table.rows:
                for cell in row.cells:
                    if ".tar" in cell.text:
                        if not any(
                            f"{patch_dir_name}.tar" in word
                            for word in cell.text.split()
                        ):
                            cell.text = cell.text.replace(
                                ".tar", f"{patch_dir_name}.tar"
                            )
                    if "월 증분 패치" in cell.text and not cell.text.startswith(
                        f"{month}월"
                    ):
                        cell.text = cell.text.replace(
                            "월 증분 패치", f"{month}월 증분 패치"
                        )
                target_numbers = {"4", "6", "7", "8", "9"}
                for row in table.rows:
                    if row.cells[0].text.strip() in target_numbers:
                        for cell in row.cells:
                            if row.cells[0].text.strip() == "9":
                                header_text = f"하위 서버 PMS 클라이언트> 최신 패치 탭에서 발표일 {current_year}년 {month}월까지 노출 확인\n\n"
                                if len(row.cells) > 1:
                                    right_cell = row.cells[1]
                                    content = right_cell.text if right_cell.text else ""

                                    content = (
                                        f"21H2({count_21h2}) - {kb_ids_21h2}\n"
                                        f"22H2({count_22h2}) - {kb_ids_22h2}\n\n"
                                        f"Office 2016({office_count}) - {office_kb_ids}\n\n"
                                    )

                                    for software in software_list:
                                        version = software_versions.get(
                                            software, "버전 정보 없음"
                                        )
                                        content += f"{software} - {version}\n"

                                    right_cell.text = header_text + content
                            if "확인 사항" in cell.text:
                                if content_to_add not in cell.text:
                                    cell.text = (
                                        cell.text.split("확인 사항")[0]
                                        + "확인 사항\n"
                                        + content_to_add
                                        + "\n"
                                        + "\n".join(cell.text.split("확인 사항")[1:])
                                    )

                for row in table.rows:
                    if row.cells[0].text.strip() == "10":
                        for cell in row.cells:
                            if "Edge 브라우저 업데이트 확인" in cell.text:
                                updated_text = re.sub(
                                    r"Edge 브라우저 업데이트 확인\(\)",
                                    f"Edge 브라우저 업데이트 확인({edge_version})",
                                    cell.text,
                                )
                                cell.text = updated_text

                for row in table.rows:
                    if row.cells[0].text.strip() == "13":
                        for cell in row.cells:
                            cell.text = re.sub(
                                r"(?<!\d)년", f"{previous_year}년", cell.text
                            )
                            cell.text = re.sub(
                                r"(?<!\d)월", f"{previous_month}월", cell.text
                            )

            self.set_font_size(document, 10)
            document.save(new_file_path)
            print(f"파일이 성공적으로 생성되었습니다: {new_file_path}")

        except Exception as e:
            print(f"오류 발생: {e}")

    def get_patch_dir(self, base_path):
        subdirs = [
            f
            for f in os.listdir(base_path)
            if os.path.isdir(os.path.join(base_path, f))
        ]
        if len(subdirs) == 1:
            return subdirs[0]
        else:
            raise Exception("data 폴더 하위에 폴더가 하나만 있어야 합니다.")

    def get_previous_year_and_month(self):
        now = datetime.now()
        if now.month == 1:
            return now.year - 1, 12
        return now.year, now.month - 1

    def update_standalone_office_patch_section(self, docx_path, base_path):
        document = Document(docx_path)
        numeric_dirs = self.file_manager.get_numeric_dirs()
        patch_dir = self.file_manager.get_patch_dir()
        standalone_path = os.path.join(base_path, patch_dir, "pms_patch", "standalone")
        ms_excel_file = os.path.join(base_path, "ms_patch_list.xlsx")
        df = pd.read_excel(ms_excel_file)

        office_patch_text_32bit = (
            "오피스 2013, 2016 32/64 Bit 설치 후 독립 설치용 오피스 패치 파일 검증\n"
            "* 패치 파일 실행 시, C:\\Users\\사용자명\\AppData\\Local\\Temp\\package 경로에 압축 해제 후 패치 설치 시도\n\n"
            "▣ 확인 사항\n"
            "A. 오피스 2016 32 Bit\n"
        )

        office_patch_text_64bit = "\nB. 오피스 2016 64 Bit\n"

        for idx, dir in enumerate(numeric_dirs, 1):
            dir_path = os.path.join(standalone_path, dir)
            exe_files = [f for f in os.listdir(dir_path) if f.endswith(".exe")]

            exe_files_32bit = []
            exe_files_64bit = []

            for exe_file in exe_files:
                file_name_no_ext = os.path.splitext(exe_file)[0]
                matched_rows = df[
                    df["패치파일"].str.contains(file_name_no_ext, na=False)
                ]

                for _, row in matched_rows.iterrows():
                    if "32비트" in str(row["제목"]):
                        exe_files_32bit.append(exe_file)
                    elif "64비트" in str(row["제목"]):
                        exe_files_64bit.append(exe_file)

            office_patch_text_32bit += f"{idx}) {dir}\n ① {dir} 폴더 선택\n ② 실행\n"
            for exe_file in exe_files_32bit:
                office_patch_text_32bit += f"    {exe_file}\n"

            office_patch_text_64bit += f"{idx}) {dir}\n ① {dir} 폴더 선택\n ② 실행\n"
            for exe_file in exe_files_64bit:
                office_patch_text_64bit += f"    {exe_file}\n"

        full_office_patch_text = office_patch_text_32bit + office_patch_text_64bit

        for table in document.tables:
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    if "독립 설치용 오피스 패치 검증" in cell.text:
                        if row_index + 1 < len(table.rows):
                            target_cell = table.rows[row_index + 1].cells[cell_index]
                            target_cell.text = full_office_patch_text
                            break

        self.set_font_size(document, 10)
        document.save(docx_path)

    def update_standalone_software_patch_section(
        self, docx_path, sw_excel_path, base_path
    ):
        document = Document(docx_path)
        df = pd.read_excel(sw_excel_path)
        patch_version_map = dict(zip(df["패치명"], df["버전"]))
        software_list = ["Adobe Acrobat Reader", "BandiZip", "Chrome"]

        adobe_version = patch_version_map.get("Adobe Acrobat Reader", "버전 정보 없음")
        bandizip_version = patch_version_map.get("BandiZip", "버전 정보 없음")
        chrome_version = patch_version_map.get("Chrome", "버전 정보 없음")

        patch_dir = self.get_patch_dir(base_path)
        standalone_path = os.path.join(base_path, patch_dir, "pms_patch", "standalone")

        software_files = {}
        for software in software_list:
            software_dir = os.path.join(standalone_path, software)
            if os.path.exists(software_dir):
                x86_files = []
                x64_files = []
                for root, dirs, files in os.walk(software_dir):
                    for file in files:
                        if "x64" in file.lower() or "64" in file:
                            x64_files.append(file)
                        elif "x86" in file.lower() or "32" in file:
                            x86_files.append(file)
                        else:
                            x86_files.append(file)
                software_files[software] = {"x86": x86_files, "x64": x64_files}
            else:
                print(f"\n[{software}] 폴더가 존재하지 않습니다.")
                software_files[software] = {"x86": [], "x64": []}

        def format_files(files):
            return (
                "\n".join([f"    {file}" for file in files])
                if files
                else "    파일 없음"
            )

        software_patch_text = (
            f"Acrobat Reader DC, BandiZip, Chrome 32/64 Bit\n"
            "설치 후 독립 설치용 일반 SW 패치 파일 검증\n\n"
            "▣ 확인 사항\n"
            f"A. Adobe Acrobat Reader ({adobe_version})\n"
            "1) Adobe Acrobat Reader 32 Bit 설치\n"
            " ① Adobe Acrobat Reader 폴더 선택\n"
            " ② 실행\n"
            f"{format_files(software_files.get('Adobe Acrobat Reader').get('x86'))}\n\n"
            f"B. BandiZip ({bandizip_version})\n"
            "1) BandiZip 32 Bit 설치\n"
            " ① BandiZip 폴더 선택\n"
            " ② 실행\n"
            f"{format_files(software_files.get('BandiZip').get('x86'))}\n"
            "2) BandiZip 64 Bit 설치\n"
            " ① BandiZip 폴더 선택\n"
            " ② 실행\n"
            f"{format_files(software_files.get('BandiZip').get('x64'))}\n\n"
            f"C. Chrome ({chrome_version})\n"
            "1) Chrome 32 Bit 설치\n"
            " ① Chrome 폴더 선택\n"
            " ② 실행\n"
            f"{format_files(software_files.get('Chrome').get('x86'))}\n"
            "2) Chrome 64 Bit 설치\n"
            " ① Chrome 폴더 선택\n"
            " ② 실행\n"
            f"{format_files(software_files.get('Chrome').get('x64'))}"
        )

        for table in document.tables:
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    if "독립 설치용 일반 SW 패치 검증" in cell.text:
                        if row_index + 1 < len(table.rows):
                            target_cell = table.rows[row_index + 1].cells[cell_index]
                            target_cell.text = software_patch_text
                            break

        self.set_font_size(document, 10)
        document.save(docx_path)

    def process_document(self):
        base_path = self.base_path
        patch_dir_name = self.file_manager.get_patch_dir()
        v1_dir_path = self.file_manager.find_v1_dir()
        ms_excel_file = os.path.join(base_path, "ms_patch_list.xlsx")
        sw_excel_file = os.path.join(base_path, "sw_patch_list.xlsx")
        docx_file_path = os.path.join(base_path, "증분 패치 검증 QA 결과서.docx")

        now = datetime.now()
        month = now.strftime("%m")
        date_str = now.strftime("%y%m%d")
        new_docx_file_path = os.path.join(
            base_path, f"{month}월 증분 패치 검증 QA 결과서_{date_str}.docx"
        )

        file_list_content = self.file_manager.list_files_in_v1(
            v1_dir_path, ms_excel_file
        )

        self.update_docx_with_content(
            docx_file_path, file_list_content, new_docx_file_path, month
        )
        self.update_standalone_office_patch_section(new_docx_file_path, base_path)
        self.update_standalone_software_patch_section(
            new_docx_file_path, sw_excel_file, base_path
        )
        self.update_completion_date(new_docx_file_path)
